"""Tests for text-hash secondary duplicate detection.

Byte-different but text-identical PDFs (e.g. re-renders with new metadata
timestamps) should be caught by ``find_duplicates``' second pass.
"""

import hashlib
from pathlib import Path
from unittest.mock import patch

import pytest

from docorganizer.cli import find_duplicates
from docorganizer.extractor import compute_text_hash


def _write_file(path: Path, content: bytes | str) -> Path:
    path.parent.mkdir(parents=True, exist_ok=True)
    if isinstance(content, str):
        path.write_text(content)
    else:
        path.write_bytes(content)
    return path


# ── compute_text_hash unit ─────────────────────────────────────────────────


class TestComputeTextHash:
    def test_returns_none_for_unsupported_extension(self, tmp_path):
        f = _write_file(tmp_path / "x.txt", "hello")
        assert compute_text_hash(f) is None

    def test_returns_none_for_missing_file(self, tmp_path):
        assert compute_text_hash(tmp_path / "missing.pdf") is None

    def test_normalizes_whitespace(self, tmp_path):
        # Two .docx with identical visible text but different whitespace
        # should hash to the same value. We exercise this through the
        # docx path since it's deterministic and doesn't need a real PDF.
        import docx
        d1 = docx.Document()
        d1.add_paragraph(
            "The quick brown fox jumps over the lazy dog. "
            "Sphinx of black quartz, judge my vow."
        )
        d1.save(tmp_path / "a.docx")
        d2 = docx.Document()
        d2.add_paragraph("The quick brown fox jumps over the lazy dog.")
        d2.add_paragraph("Sphinx of black quartz, judge my vow.")
        d2.save(tmp_path / "b.docx")

        h1 = compute_text_hash(tmp_path / "a.docx")
        h2 = compute_text_hash(tmp_path / "b.docx")
        assert h1 is not None
        assert h1 == h2

    def test_different_text_different_hash(self, tmp_path):
        import docx
        d1 = docx.Document()
        d1.add_paragraph("Document number one with enough content to hash safely.")
        d1.save(tmp_path / "a.docx")
        d2 = docx.Document()
        d2.add_paragraph("Document number two with enough content to hash safely.")
        d2.save(tmp_path / "b.docx")
        assert compute_text_hash(tmp_path / "a.docx") != compute_text_hash(tmp_path / "b.docx")

    def test_too_short_returns_none(self, tmp_path):
        import docx
        d = docx.Document()
        d.add_paragraph("hi")  # below MIN_TEXT_HASH_LENGTH
        d.save(tmp_path / "tiny.docx")
        assert compute_text_hash(tmp_path / "tiny.docx") is None


# ── find_duplicates integration ────────────────────────────────────────────


class TestTextHashSecondPass:
    def test_byte_identical_still_caught_first(self, test_ctx):
        """Byte-identical files are caught by the first pass and never need text extraction."""
        archived = test_ctx.root_folder / "Germany" / "Steuer" / "filed.pdf"
        _write_file(archived, "identical bytes content")
        inbox_file = _write_file(test_ctx.inbox / "incoming.pdf", "identical bytes content")

        duplicates, remaining = find_duplicates([inbox_file], test_ctx)
        assert len(duplicates) == 1
        assert duplicates[0].match_type == "bytes"
        assert remaining == []

    def test_text_identical_byte_different_caught(self, test_ctx, tmp_path):
        """When text-hash returns the same value but bytes differ, treat as duplicate."""
        archived = test_ctx.root_folder / "Germany" / "Steuer" / "filed.pdf"
        archived.parent.mkdir(parents=True, exist_ok=True)
        archived.write_bytes(b"version-A-bytes")
        inbox_file = test_ctx.inbox / "rerendered.pdf"
        inbox_file.write_bytes(b"version-B-bytes-different-length")

        # Patch compute_text_hash to return identical hashes for both files,
        # simulating a re-render with metadata-only differences.
        sentinel = "deadbeef" * 8

        def fake_text_hash(path):
            if path in (archived, inbox_file):
                return sentinel
            return None

        with patch("docorganizer.cli.compute_text_hash", side_effect=fake_text_hash):
            duplicates, remaining = find_duplicates([inbox_file], test_ctx)

        assert len(duplicates) == 1
        assert duplicates[0].match_type == "text"
        assert duplicates[0].existing_path == archived
        assert duplicates[0].file_hash == sentinel
        assert remaining == []

    def test_no_text_hash_means_no_false_positive(self, test_ctx):
        """When text extraction is unavailable for both files, do not dedupe."""
        archived = test_ctx.root_folder / "Germany" / "Steuer" / "old.pdf"
        archived.parent.mkdir(parents=True, exist_ok=True)
        archived.write_bytes(b"some bytes")
        inbox_file = test_ctx.inbox / "new.pdf"
        inbox_file.write_bytes(b"different bytes")

        # compute_text_hash returning None for everything (e.g. extraction failed)
        with patch("docorganizer.cli.compute_text_hash", return_value=None):
            duplicates, remaining = find_duplicates([inbox_file], test_ctx)

        assert duplicates == []
        assert remaining == [inbox_file]

    def test_text_dedup_within_same_batch(self, test_ctx):
        """Two inbox files with same text-hash but different bytes — second is a dup."""
        a = test_ctx.inbox / "first.pdf"
        b = test_ctx.inbox / "second.pdf"
        a.write_bytes(b"different-bytes-A")
        b.write_bytes(b"different-bytes-B-longer")

        sentinel = "abcdef00" * 8

        def fake_text_hash(path):
            if path in (a, b):
                return sentinel
            return None

        with patch("docorganizer.cli.compute_text_hash", side_effect=fake_text_hash):
            duplicates, remaining = find_duplicates([a, b], test_ctx)

        assert len(duplicates) == 1
        assert duplicates[0].inbox_path == b
        assert duplicates[0].existing_path == a
        assert duplicates[0].match_type == "text"
        assert remaining == [a]

    def test_archive_text_hash_not_computed_when_not_needed(self, test_ctx):
        """If every inbox file is a byte-duplicate, archive text-hashes are never computed."""
        archived = test_ctx.root_folder / "Germany" / "Steuer" / "filed.pdf"
        _write_file(archived, "shared content here yep")
        inbox_file = _write_file(test_ctx.inbox / "x.pdf", "shared content here yep")

        with patch("docorganizer.cli.compute_text_hash") as fake:
            duplicates, remaining = find_duplicates([inbox_file], test_ctx)

        assert len(duplicates) == 1
        assert duplicates[0].match_type == "bytes"
        # First pass found the match — text-hash never invoked
        fake.assert_not_called()
